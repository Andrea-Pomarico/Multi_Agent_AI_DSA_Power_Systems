"""
Documentation generator for the Multi-Agent DSA Pipeline.
Produces a well-formatted Word (.docx) file using python-docx.
"""

import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False
    print("python-docx not found. Install with: pip install python-docx")
    sys.exit(1)


OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Multi_Agent_DSA_Pipeline_Documentation.docx")


def set_heading_color(paragraph, r, g, b):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def add_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def add_table_row(table, cells, bold=False, shade=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        if bold:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        if shade:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), shade)
            tcPr.append(shd)
    return row


def style_table(table):
    table.style = 'Table Grid'
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)


def build_doc():
    doc = Document()

    # --- Page margins ---
    from docx.oxml.ns import nsmap
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    # =========================================================
    # TITLE PAGE
    # =========================================================
    doc.add_picture  # no-op, just checking
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Multi-Agent Dynamic Security Assessment Pipeline")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Technical Documentation — Architecture, Implementation & Usage Guide")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x26, 0x59, 0x9E)
    sub_run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run("Andrea Pomarico")

    doc.add_page_break()

    # =========================================================
    # TABLE OF CONTENTS
    # =========================================================
    h = doc.add_heading("Table of Contents", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)
    add_toc(doc)
    doc.add_page_break()

    # =========================================================
    # 1. EXECUTIVE OVERVIEW
    # =========================================================
    h = doc.add_heading("1. Executive Overview", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)

    doc.add_paragraph(
        "The Multi-Agent Dynamic Security Assessment (DSA) Pipeline is an integrated software framework "
        "designed to automate the full workflow of power systems transient stability analysis. The system "
        "combines a high-fidelity physics simulation engine (DIgSILENT PowerFactory) with Large Language "
        "Model (LLM) intelligence (Google Gemini / Groq) to produce publication-quality engineering reports, "
        "PowerPoint presentations, and actionable grid mitigation recommendations — with minimal user intervention."
    )

    doc.add_paragraph(
        "The pipeline is composed of ten specialised agents, each responsible for one stage of the analysis. "
        "Agents communicate through structured data objects and plain-text artefacts: simulation CSV files, "
        "numerical KPI dictionaries, plot images (PNG), Word documents (.docx) and PowerPoint decks (.pptx)."
    )

    h2 = doc.add_heading("Key Capabilities", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)
    for item in [
        "Automated end-to-end RMS transient stability simulation via DIgSILENT PowerFactory Python API.",
        "Natural-language or JSON-based simulation scenario specification.",
        "Multi-case batch execution with automatic case comparison.",
        "Numerical KPI extraction: voltage nadir, rotor angle excursion, speed settling time, and more.",
        "AI-generated narrative engineering reports (draft → review → final).",
        "Automatic grid topology visualisation and network graph export.",
        "AI-powered mitigation recommendations ranked by urgency and complexity.",
        "Word document and PowerPoint generation with embedded plots.",
        "Provider-agnostic LLM abstraction: switch between Gemini and Groq via configuration.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    doc.add_page_break()

    # =========================================================
    # 2. SYSTEM ARCHITECTURE
    # =========================================================
    h = doc.add_heading("2. System Architecture", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)

    doc.add_paragraph(
        "The pipeline is structured as a strictly sequential chain of ten agents. Each agent consumes the "
        "outputs of the previous stage and produces artefacts consumed by the next. The two entry points "
        "are run_rms_pipeline() for a single simulation case, and run_rms_multi_case_pipeline() for "
        "batch execution across multiple fault scenarios."
    )

    h2 = doc.add_heading("2.1 Agent Overview Table", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(["Agent", "Name", "Purpose", "Key Input", "Key Output"]):
        hdr[i].text = txt
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)

    rows_data = [
        ("0", "Intake Agent", "Parse user input; produce case list", "Natural language / JSON / Defaults", "SimulationConfig + cases list"),
        ("1", "Simulation Agent", "Execute RMS simulation in PowerFactory", "SimulationConfig", "CSV time-series + network graph"),
        ("2", "Analysis Agent", "Extract numerical KPIs from CSV", "CSV results + timing params", "Numerics dict + stats text"),
        ("3", "Plot Agent", "Generate time-series visualisations", "Numerics dict + config", "PNG plot files"),
        ("4", "Summary Agent (LLM)", "Draft narrative stability report", "KPI text block + study case", "Draft report text"),
        ("5", "Review Agent (LLM)", "QA check draft vs raw KPIs", "KPI block + draft report", "Numbered correction list"),
        ("6", "Final Report Agent (LLM)", "Apply corrections, polish report", "Draft + corrections", "Final polished report"),
        ("7", "Comparison Agent (LLM)", "Multi-case comparative analysis", "All case KPIs + case names", "Comparative report"),
        ("8", "Presentation Agent", "Build PowerPoint deck", "Report text + PNG plots", "PPTX file"),
        ("9", "Mitigation Agent (LLM)", "Propose ranked grid improvements", "Report + KPIs + grid topology", "Mitigation action plan"),
    ]

    fills = ['F2F2F2', 'FFFFFF']
    for idx, row_d in enumerate(rows_data):
        row = tbl.add_row()
        for i, text in enumerate(row_d):
            row.cells[i].text = text
            tc = row.cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fills[idx % 2])
            tcPr.append(shd)

    doc.add_paragraph()

    h2 = doc.add_heading("2.2 Single-Case Data Flow", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    doc.add_paragraph(
        "The following describes the data flow for a single-case execution (run_rms_pipeline):"
    )

    flow_steps = [
        ("Step 1 — Intake (Agent 0)", "User provides a fault scenario via natural language, JSON, or defaults. "
         "Agent 0 (LLM-powered) parses this into a validated SimulationConfig dataclass, including fault type, "
         "fault element, switching times, and signal list."),
        ("Step 2 — Simulation (Agent 1)", "DIgSILENTAgent connects to PowerFactory, activates the project and "
         "study case, exports the network topology graph, runs a pre-fault load flow, applies the fault event "
         "(EvtShc or EvtSwitch), executes the RMS simulation (ComInc + ComSim), and exports results to CSV via ComRes."),
        ("Step 3 — Analysis (Agent 2)", "The CSV is parsed (semicolon-delimited, two-row headers). Per-signal KPIs "
         "are computed: voltage nadir, post-fault mean/std, settling time, limit violation flag; generator speed "
         "min/max/settle; rotor angle excursion and post-fault oscillation metrics."),
        ("Step 4 — Plotting (Agent 3)", "Matplotlib generates voltage and speed time-series plots with fault-event "
         "markers, ±10% bands, and per-violated-bus close-up plots. All saved as PNG files."),
        ("Step 5 — LLM Reporting (Agents 4, 5, 6)", "Agent 4 (MODEL_FAST) produces a structured draft report "
         "from the KPI text block. Agent 5 (MODEL_FAST) cross-checks the draft against raw KPIs and returns a "
         "correction list. Agent 6 (MODEL_SMART) applies corrections and produces the final polished report."),
        ("Step 6 — Presentation (Agent 8)", "python-pptx builds a deck: title slide, report text chunked to "
         "12 lines per slide, followed by one full-page slide per PNG plot."),
        ("Step 7 — Mitigation (Agent 9)", "A two-stage LLM call: first, grid topology vulnerabilities are "
         "identified from the structured grid_data dict; second, ranked mitigation actions (immediate, "
         "short-term, long-term) are generated."),
        ("Step 8 — Artefact Persistence", "report_utils saves: pipeline log CSV, KPI CSV, LLM summary CSV, "
         "plain-text report, Word document (.docx with embedded plots), and the PowerPoint file."),
    ]

    for title, desc in flow_steps:
        p = doc.add_paragraph()
        run = p.add_run(title + ": ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        p.add_run(desc)

    h2 = doc.add_heading("2.3 Multi-Case Data Flow", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    doc.add_paragraph(
        "run_rms_multi_case_pipeline() iterates through a cases list (from the 'cases' or 'faults' array "
        "in simulation_config.json, or from Agent 0 natural-language parsing). Each case is executed "
        "independently through the single-case pipeline. After all cases complete:"
    )
    for item in [
        "Agent 7 (Comparison Agent) performs a cross-case LLM analysis of all KPI blocks, identifying the "
        "most severe scenarios and trends.",
        "Agent 3 generates multi-case overlay plots: one voltage overlay per violated bus, one speed overlay "
        "for the top-N most deviated generators.",
        "A consolidated Word report and PowerPoint deck are generated for the full batch.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    doc.add_page_break()

    # =========================================================
    # 3. FILE STRUCTURE
    # =========================================================
    h = doc.add_heading("3. File Structure & Module Descriptions", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)

    modules = [
        (
            "Multi_Agent_RMS_google.py",
            "Main Orchestrator",
            [
                "run_rms_pipeline(cfg, emit_final_report, emit_final_presentation): "
                "Executes the complete single-case pipeline. Creates a timestamped output folder, "
                "calls all agents in sequence, handles errors, logs step durations, and returns a "
                "comprehensive report dictionary.",
                "run_rms_multi_case_pipeline(cases_list, base_cfg): "
                "Iterates over a list of case specifications, runs each through the single-case "
                "pipeline, then calls comparison and overlay plot agents on the aggregated results.",
                "_build_case_config(base_cfg, case_spec, parent_dir): "
                "Maps a case JSON dict onto a SimulationConfig, normalising fault_type aliases "
                "('generator' → 'gen_switch') and switch state strings ('open' → 0, 'close' → 1).",
                "_cases_from_config_json(cfg_path): "
                "Extracts and validates the 'cases' or 'faults' array from the configuration JSON.",
                "Entry point: loads agent_prompts.csv, configures the LLM client, reads simulation_config.json, "
                "detects single- vs. multi-case mode, and launches the appropriate pipeline.",
            ]
        ),
        (
            "Agent_DIgSILENT.py",
            "DIgSILENT PowerFactory Integration",
            [
                "SimulationConfig (dataclass): Complete parameter set for one RMS simulation, including "
                "project path, study case names, fault type and element, event timings, integration "
                "step size, output directory, signal list, and flags to enable optional agents.",
                "DIgSILENTAgent: Wrapper around the PowerFactory Python API. Manages application and "
                "project lifecycle. Key methods: connect(), activate_study_case(), export_grid_graph(), "
                "run_loadflow(), run_rms_simulation(), _apply_fault_event(), export_results_to_csv(), "
                "run_pipeline().",
                "export_grid_graph(): Collects bus (voltage, angle, Ikss), line (loading %, power flow), "
                "load, and generator data from the active load-flow solution. Draws a NetworkX topology "
                "diagram (colour-coded by loading), exports it as PNG, and returns a structured grid_data "
                "dict used by Agent 9.",
                "_apply_fault_event(): Creates EvtShc (3-phase bus fault), EvtSwitch (line trip or "
                "generator disconnection) events in the Fault.IntEvt event folder based on fault_type.",
                "Logger: Timestamped console logger with status icons (✅ success, ❌ error, ⚠️ warning).",
            ]
        ),
        (
            "agent_intake.py",
            "Agent 0 — Intake Agent",
            [
                "intake_agent(cfg): Presents the user with three input modes: (1) natural language description "
                "parsed by LLM into a case list, (2) JSON or key:value overrides for a single case, "
                "(3) empty input to use configuration defaults.",
                "parse_natural_language_input(user_text): Calls the LLM with the agent_0_input_parser "
                "system prompt. Extracts and validates the JSON case array from the LLM response.",
                "_coerce_user_request(raw): Flexible parser that attempts JSON parsing first, then "
                "falls back to line-by-line key:value parsing with type coercion.",
            ]
        ),
        (
            "agent_simulation.py",
            "Agent 1 — Simulation Agent",
            [
                "simulation_agent(cfg): Thin wrapper that instantiates DIgSILENTAgent and calls "
                "run_pipeline(). Returns the PowerFactory report dict containing CSV path, "
                "grid_data, and execution status.",
            ]
        ),
        (
            "agent_analysis.py",
            "Agent 2 — Analysis Agent",
            [
                "analysis_agent(csv_path, t_fault, t_clear, fault_type, ...): Main KPI extraction "
                "function. Parses the PowerFactory CSV, classifies signals by variable name, and "
                "computes per-signal statistics across pre-fault, fault, and post-fault time windows.",
                "Voltage KPIs: pre_mean, nadir (fault-period minimum), post_mean, post_std, "
                "settle_s, v_t0, v_t_end, post_clear_min/max with timestamps, post_min/max "
                "overall, final_val, out_of_limit_after_check (0.9–1.1 p.u. band violation flag).",
                "Speed KPIs: pre_mean, min, max, post_mean, post_std, settle_s (±0.5% of 1 p.u.).",
                "Angle KPIs: pre_mean, angle_t0, angle_t_end, min, max, delta_max, delta_span, post_std.",
                "_settling_time(time, arr, ref, band, t_start): Finds the first time after t_start "
                "where the signal remains within ref ± band for the rest of the window.",
                "_is_switched_out_generator(): Excludes generators that are the target of a "
                "gen_switch open event from speed and angle analysis.",
            ]
        ),
        (
            "agent_plot.py",
            "Agent 3 — Plot Agent",
            [
                "plot_agent(numerics, cfg): Generates single-case plots using matplotlib (Agg backend). "
                "Produces: a voltage time-series PNG for all buses with ±10% limit lines, a speed "
                "time-series PNG for all online generators, and individual close-up PNGs for each "
                "bus with a post-fault voltage violation.",
                "plot_voltage_case_comparison(): Multi-case overlay plot — one PNG per violated bus, "
                "with one coloured trace per case.",
                "plot_speed_case_comparison(): Multi-case overlay plot for the top-N generators "
                "(ranked by worst |speed − 1.0| deviation across all cases).",
            ]
        ),
        (
            "agent_llm_reporting.py",
            "Agents 4, 5, 6, 7 — LLM Reporting",
            [
                "_build_kpi_block(numerics, compact, ...): Formats the numerics dict as a "
                "structured text block for LLM consumption. Compact mode ranks signals by "
                "severity and returns only the top-N worst.",
                "summary_agent(numerics, study_case): Agent 4. Uses MODEL_FAST. Generates "
                "a structured draft report covering voltage stability, rotor angle/frequency "
                "stability, and an overall system verdict.",
                "review_agent(numerics, summary): Agent 5. Uses MODEL_FAST. Cross-checks the "
                "draft report against raw KPIs and returns a numbered correction list.",
                "final_report_agent(summary, improvements, study_case): Agent 6. Uses MODEL_SMART. "
                "Applies all reviewer corrections and returns a polished final report.",
                "comparison_agent(results_dict, case_names): Agent 7. Uses MODEL_SMART. Compares "
                "KPIs and narrative across multiple cases, identifying the worst-case scenario and trends.",
            ]
        ),
        (
            "agent_presentation.py",
            "Agent 8 — Presentation Agent",
            [
                "presentation_agent(...): Requires python-pptx. Builds a slide deck: Slide 1 is the "
                "title with study case and timestamp; subsequent slides contain the report text "
                "(12 lines per slide); optional source-document reference slide; one full-page image "
                "slide per PNG plot. Returns (ok, pptx_path, message).",
            ]
        ),
        (
            "agent_mitigation.py",
            "Agent 9 — Mitigation Agent",
            [
                "mitigation_agent(final_report, numerics, grid_data, study_case): Two-stage LLM call. "
                "Stage A: submits the structured grid_data (buses, lines, loads, generators with "
                "pre-fault load-flow values) and asks the LLM to identify vulnerabilities "
                "(voltage deviations, overloaded branches, weak short-circuit buses). "
                "Stage B: combines Stage A findings with the final report and KPIs and requests "
                "a ranked action plan with IMMEDIATE, SHORT-TERM, and LONG-TERM sections. "
                "Each action includes Type, Target, Issue, Mechanism, Impact, and Complexity fields.",
            ]
        ),
        (
            "llm_client.py",
            "LLM Provider Abstraction",
            [
                "configure_llm_from_config(raw_cfg): Reads provider, API keys, and model aliases "
                "from the configuration JSON. Supports nested structure or flat legacy keys. Also "
                "checks GEMINI_API_KEY, GROQ_API_KEY, LLM_PROVIDER environment variables.",
                "run_agent(system_prompt, user_message, max_tokens, model): Routes the call to "
                "Gemini (google-genai SDK) or Groq (groq SDK) at temperature=0.1. MODEL_FAST and "
                "MODEL_SMART are symbolic aliases resolved at call time.",
                "run_vision_agent(..., image_path): Gemini supports full multimodal input (embeds "
                "PNG/JPG). Groq silently falls back to text-only if no vision-capable model is active.",
                "get_llm_runtime_info(): Returns current provider name and resolved model strings for "
                "fast and smart tiers.",
            ]
        ),
        (
            "prompt_loader.py",
            "System Prompt Manager",
            [
                "load_prompts(csv_path): Loads agent system prompts from agent_prompts.csv "
                "(columns: agent_id, agent_name, system_prompt). Results cached in a module-level dict.",
                "get_prompt(agent_id, default): Retrieves a cached prompt by agent ID.",
            ]
        ),
        (
            "report_utils.py",
            "Artefact Persistence",
            [
                "save_pipeline_log(log_rows, out_dir, label): Saves execution timeline as CSV "
                "(timestamp | step | status | message | duration_s).",
                "save_kpi_csv(numerics, out_dir, label): Tidy CSV — one row per "
                "(signal_type, object, metric, value, unit). Angles converted to degrees.",
                "save_summary_csv(summary, cfg, out_dir, label): Run metadata + LLM summary lines.",
                "save_final_report_txt(improvements, final_report, out_dir, label): Plain-text report.",
                "save_report_docx(final_report, out_dir, label, plot_paths): Custom OOXML writer "
                "(no python-docx dependency). Uses stdlib zipfile + raw XML strings to embed "
                "PNG images as inline drawings in a fully compliant .docx file.",
            ]
        ),
    ]

    for filename, role, bullets in modules:
        h2 = doc.add_heading(filename, level=2)
        set_heading_color(h2, 0x26, 0x59, 0x9E)
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(f"Role: {role}")
        role_run.bold = True
        role_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        for b in bullets:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(b)

    doc.add_page_break()

    # =========================================================
    # 4. CONFIGURATION
    # =========================================================
    h = doc.add_heading("4. Configuration Reference", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)

    doc.add_paragraph(
        "All pipeline behaviour is controlled through simulation_config.json in the project directory. "
        "The file has two top-level sections: an 'llm' block for API and model settings, and simulation "
        "parameters for the PowerFactory run."
    )

    h2 = doc.add_heading("4.1 LLM Configuration Block", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = 'Table Grid'
    for i, txt in enumerate(["Key", "Type / Values", "Description"]):
        tbl2.rows[0].cells[i].text = txt
        for para in tbl2.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        tc = tbl2.rows[0].cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)

    llm_params = [
        ("llm.provider", "\"gemini\" | \"groq\"", "Active LLM provider"),
        ("llm.api_keys.gemini", "string", "Google AI Studio API key"),
        ("llm.api_keys.groq", "string", "Groq API key"),
        ("llm.models_by_provider.gemini.fast", "string", "Gemini model for fast tier (e.g. gemini-2.5-flash-lite)"),
        ("llm.models_by_provider.gemini.smart", "string", "Gemini model for smart tier (e.g. gemini-2.5-flash)"),
        ("llm.models_by_provider.groq.fast", "string", "Groq model for fast tier (e.g. llama-3.1-8b-instant)"),
        ("llm.models_by_provider.groq.smart", "string", "Groq model for smart tier (e.g. llama-3.3-70b-versatile)"),
    ]
    fills = ['F2F2F2', 'FFFFFF']
    for idx, (k, t, d) in enumerate(llm_params):
        row = tbl2.add_row()
        row.cells[0].text = k
        row.cells[1].text = t
        row.cells[2].text = d
        for ci in range(3):
            tc = row.cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fills[idx % 2])
            tcPr.append(shd)

    doc.add_paragraph()

    h2 = doc.add_heading("4.2 Simulation Parameters", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = 'Table Grid'
    for i, txt in enumerate(["Parameter", "Type / Values", "Description"]):
        tbl3.rows[0].cells[i].text = txt
        for para in tbl3.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        tc = tbl3.rows[0].cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)

    sim_params = [
        ("project_path", "string", "PowerFactory project path (UNC or absolute)"),
        ("study_case", "string", "Target study case name within the project"),
        ("base_study_case", "string", "Fallback study case used if target is not found"),
        ("output_dir", "string", "Directory where all artefacts are written"),
        ("run_label", "string", "Prefix for all output file names"),
        ("result_name", "string", "PowerFactory result object name (e.g. All calculations.ElmRes)"),
        ("word_document", "0 | 1", "Generate a per-case Word document (multi-case mode)"),
        ("final_word_document", "0 | 1", "Generate the consolidated final Word document"),
        ("final_presentation", "0 | 1", "Generate the consolidated PowerPoint presentation"),
        ("run_review_agent", "0 | 1", "Enable Agent 5 (review / QA) in LLM reporting"),
        ("run_final_report_agent", "0 | 1", "Enable Agent 6 (final report polish)"),
        ("run_mitigation_agent", "0 | 1", "Enable Agent 9 (grid mitigation recommendations)"),
        ("cases", "array", "List of fault scenario objects (see Section 4.3)"),
    ]
    for idx, (k, t, d) in enumerate(sim_params):
        row = tbl3.add_row()
        row.cells[0].text = k
        row.cells[1].text = t
        row.cells[2].text = d
        for ci in range(3):
            tc = row.cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fills[idx % 2])
            tcPr.append(shd)

    doc.add_paragraph()

    h2 = doc.add_heading("4.3 Case Specification Object", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    doc.add_paragraph(
        "Each element of the 'cases' array describes one simulation scenario. The fault_type "
        "field determines which additional fields are required:"
    )

    tbl4 = doc.add_table(rows=1, cols=4)
    tbl4.style = 'Table Grid'
    for i, txt in enumerate(["Field", "Type", "Required For", "Description"]):
        tbl4.rows[0].cells[i].text = txt
        for para in tbl4.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        tc = tbl4.rows[0].cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)

    case_params = [
        ("case_name", "string", "All", "Label used in output filenames and reports"),
        ("fault_type", "\"bus\" | \"line\" | \"gen_switch\"", "All", "Type of fault event to simulate"),
        ("fault_element", "string", "All", "PowerFactory object name of faulted element (e.g. Bus 25.ElmTerm)"),
        ("switch_element", "string", "gen_switch", "Generator or breaker object to switch"),
        ("switch_state", "\"open\" | \"close\" | 0 | 1", "gen_switch", "Target switch state (0=open, 1=close)"),
        ("t_switch", "float (s)", "gen_switch", "Time at which the switch event occurs"),
        ("t_start", "float (s)", "All", "Simulation start time (default: 0.0)"),
        ("t_fault", "float (s)", "bus / line", "Time of fault application"),
        ("t_clear", "float (s)", "bus / line", "Time of fault clearing"),
        ("t_end", "float (s)", "All", "Simulation end time"),
        ("dt_rms", "float (s)", "All", "RMS integration time step (default: 0.01)"),
        ("word_document", "0 | 1", "All", "Override per-case Word document flag"),
    ]
    for idx, row_d in enumerate(case_params):
        row = tbl4.add_row()
        for i, text in enumerate(row_d):
            row.cells[i].text = text
            tc = row.cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fills[idx % 2])
            tcPr.append(shd)

    doc.add_page_break()

    # =========================================================
    # 5. DATA STRUCTURES
    # =========================================================
    h = doc.add_heading("5. Key Data Structures", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)

    h2 = doc.add_heading("5.1 SimulationConfig", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)
    doc.add_paragraph(
        "SimulationConfig is a Python dataclass defined in Agent_DIgSILENT.py. It is the central "
        "parameter object passed between Agent 0, Agent 1, and the orchestrator. Key fields include:"
    )
    for item in [
        "project_path, study_case, base_study_case — PowerFactory project and case identification.",
        "fault_type ('bus' | 'line' | 'gen_switch'), fault_element, switch_element, switch_state, t_switch — fault scenario definition.",
        "t_start, t_fault, t_clear, t_end, dt_rms — simulation time window and integration step.",
        "output_dir, run_label, result_name — artefact storage and naming.",
        "signals — list of (object_name, variable_name, friendly_label) tuples specifying which PowerFactory signals to export.",
        "run_review_agent, run_final_report_agent, run_mitigation_agent — integer flags to enable/disable optional pipeline stages.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    h2 = doc.add_heading("5.2 Numerics Dictionary", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)
    doc.add_paragraph(
        "The numerics dict is produced by Agent 2 and consumed by Agents 3, 4, 5, 6, 7, and 9. "
        "It has the following top-level structure:"
    )
    p = doc.add_paragraph(style='No Spacing')
    p.add_run(
        "{ 't_fault': float, 't_clear': float, 't_end': float, 'dt': float, 'check_time': float,\n"
        "  'voltages': { bus_name: { pre_mean, nadir, post_mean, post_std, settle_s, v_t0, v_t_end,\n"
        "                             post_clear_min/max (+ timestamps), post_min/max, final_val,\n"
        "                             out_of_limit_after_check } },\n"
        "  'speeds':   { gen_name: { pre_mean, min, max, post_mean, post_std, settle_s } },\n"
        "  'angles':   { gen_name: { pre_mean, angle_t0, angle_t_end, min, max,\n"
        "                             delta_max, delta_span, post_std } },\n"
        "  'scenario': { fault_type, fault_element, switch_element, switch_state, t_switch,\n"
        "                excluded_generators },\n"
        "  '_parsed':  { 'time': np.array, 'signals': dict }  # raw CSV data }"
    ).font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)

    h2 = doc.add_heading("5.3 grid_data Dictionary", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)
    doc.add_paragraph(
        "grid_data is produced by DIgSILENTAgent.export_grid_graph() and consumed by Agent 9. "
        "It captures the pre-fault steady-state network topology:"
    )
    for item in [
        "buses: dict keyed by normalised bus name → { name (PowerFactory object name), u (voltage in p.u.), phiu (angle in degrees), ikss (short-circuit current in kA) }.",
        "lines: list of { name, bus1, bus2, loading (%), p_flow (MW) }.",
        "loads: list of { name, bus, p_ini (MW), q_ini (Mvar) }.",
        "generators: list of { name, bus, p_ini (MW), q_ini (Mvar) }.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    doc.add_page_break()

    # =========================================================
    # 6. EXTERNAL INTEGRATIONS
    # =========================================================
    h = doc.add_heading("6. External Integrations", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)

    h2 = doc.add_heading("6.1 DIgSILENT PowerFactory", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)
    doc.add_paragraph(
        "PowerFactory is accessed via its embedded Python API, available from PowerFactory 2025 SP1 "
        "at the path configured in Agent_DIgSILENT.py "
        "(C:\\Program Files\\DIgSILENT\\PowerFactory 2025 SP1\\Python\\3.13). "
        "The pipeline uses the following PowerFactory command and element objects:"
    )
    pf_objects = [
        ("ComLdf", "Load flow calculation command"),
        ("ComInc", "RMS simulation initialisation command"),
        ("ComSim", "RMS simulation execution command"),
        ("ComRes", "Result export command (to CSV)"),
        ("ComShc", "Short-circuit calculation command"),
        ("EvtShc", "Short-circuit fault event"),
        ("EvtSwitch", "Switch / breaker event"),
        ("ElmTerm", "Bus / terminal element"),
        ("ElmLne", "Transmission line element"),
        ("ElmTr2", "Two-winding transformer element"),
        ("ElmSym", "Synchronous generator element"),
        ("ElmLod", "Load element"),
    ]
    tbl_pf = doc.add_table(rows=1, cols=2)
    tbl_pf.style = 'Table Grid'
    for i, txt in enumerate(["PowerFactory Class", "Purpose"]):
        tbl_pf.rows[0].cells[i].text = txt
        for para in tbl_pf.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        tc = tbl_pf.rows[0].cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
    for idx, (cls, desc) in enumerate(pf_objects):
        row = tbl_pf.add_row()
        row.cells[0].text = cls
        row.cells[1].text = desc
        for ci in range(2):
            tc = row.cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fills[idx % 2])
            tcPr.append(shd)

    doc.add_paragraph()

    h2 = doc.add_heading("6.2 Google Gemini API", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)
    doc.add_paragraph(
        "Gemini is the default LLM provider. The google-genai Python SDK is used at temperature=0.1 "
        "to minimise hallucination while retaining fluent language. Two model tiers are configured:"
    )
    for item in [
        "MODEL_FAST (gemini-2.5-flash-lite): Used for Agents 4 and 5 where speed and low cost are prioritised.",
        "MODEL_SMART (gemini-2.5-flash): Used for Agents 6, 7, and 9 where output quality is critical.",
        "Vision support: run_vision_agent() can embed PNG images directly into the multimodal context.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    h2 = doc.add_heading("6.3 Groq API", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)
    doc.add_paragraph(
        "Groq is an alternative LLM provider selectable via the 'provider' configuration key. "
        "Fast inference is provided by Meta Llama models hosted on Groq hardware:"
    )
    for item in [
        "MODEL_FAST (llama-3.1-8b-instant): Fast tier — Agents 4, 5.",
        "MODEL_SMART (llama-3.3-70b-versatile): Smart tier — Agents 6, 7, 9.",
        "Limitation: Groq does not support vision input on default models; image arguments are silently ignored.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    doc.add_page_break()

    # =========================================================
    # 7. INPUT / OUTPUT FORMATS
    # =========================================================
    h = doc.add_heading("7. Input and Output Formats", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)

    h2 = doc.add_heading("7.1 Input Modes", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    io_modes = [
        ("Natural Language",
         "User types a description such as: 'Simulate a 3-phase bus fault at Bus 25, clearing at 80 ms, "
         "and a line trip on Line 03-04 clearing at 100 ms.' Agent 0 passes this to the LLM using the "
         "agent_0_input_parser system prompt and extracts a JSON case array."),
        ("JSON / Key:Value",
         "User pastes a JSON object or key:value lines directly at the intake prompt. The _coerce_user_request "
         "function parses and type-coerces the values. This overrides individual SimulationConfig fields for "
         "a single case."),
        ("Configuration File",
         "simulation_config.json is loaded at startup. If it contains a 'cases' or 'faults' array, the pipeline "
         "runs in multi-case mode without prompting the user. If no array is present and the user presses Enter "
         "at the intake prompt, the top-level defaults are used as a single case."),
    ]
    for title, desc in io_modes:
        p = doc.add_paragraph()
        run = p.add_run(title + ": ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        p.add_run(desc)

    h2 = doc.add_heading("7.2 Output Artefacts", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    tbl_out = doc.add_table(rows=1, cols=3)
    tbl_out.style = 'Table Grid'
    for i, txt in enumerate(["File", "Format", "Contents"]):
        tbl_out.rows[0].cells[i].text = txt
        for para in tbl_out.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        tc = tbl_out.rows[0].cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)

    out_files = [
        ("{label}_results.csv", "CSV (semicolon)", "PowerFactory time-series: time | bus voltages | gen speeds | gen angles"),
        ("{label}_kpi_summary.csv", "Tidy CSV", "One row per (signal_type, object, metric, value, unit)"),
        ("{label}_pipeline_log.csv", "Semicolon CSV", "Execution timeline: timestamp | step | status | message | duration_s"),
        ("{label}_llm_summary.csv", "Semicolon CSV", "Run metadata + LLM summary text lines"),
        ("{label}_voltages.png", "PNG", "Voltage time-series for all buses with ±10% limit bands"),
        ("{label}_speeds.png", "PNG", "Generator speed time-series for all online generators"),
        ("{label}_{bus}_voltage_violation.png", "PNG", "Close-up voltage plot for each bus with a post-fault violation"),
        ("{label}_grid_topology.png", "PNG", "NetworkX network topology diagram, colour-coded by line loading"),
        ("{label}_full_report.docx", "Word (.docx)", "Final polished report with embedded plots and mitigation section"),
        ("{label}_presentation.pptx", "PowerPoint (.pptx)", "Slide deck: title → report text → plot images"),
    ]
    for idx, row_d in enumerate(out_files):
        row = tbl_out.add_row()
        for i, text in enumerate(row_d):
            row.cells[i].text = text
            tc = row.cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fills[idx % 2])
            tcPr.append(shd)

    doc.add_page_break()

    # =========================================================
    # 8. LLM PROMPTING STRATEGY
    # =========================================================
    h = doc.add_heading("8. LLM Prompting Strategy", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)

    doc.add_paragraph(
        "All system prompts are stored in agent_prompts.csv and loaded once at startup via prompt_loader.py. "
        "This decouples prompt engineering from code, allowing prompts to be updated without touching "
        "Python source files. The CSV uses agent_id as the key (e.g. 'agent_0_input_parser', 'agent_4_summary')."
    )

    prompt_agents = [
        ("agent_0_input_parser", "Agent 0", "Converts free-form natural language into a strictly formatted "
         "JSON array of case objects. Enforces required fields per fault_type. Returns only valid JSON, "
         "no explanatory text."),
        ("agent_4_summary", "Agent 4", "Instructs the model to write a structured engineering report with "
         "sections: Executive Summary, Critical Event Timeline, Voltage Stability, Rotor Angle/Frequency "
         "Stability, System-Wide Classification (STABLE | CRITICALLY STABLE | UNSTABLE), Key Findings, "
         "and Data Gaps. Tone is factual and concise."),
        ("agent_5_review", "Agent 5", "Cross-checks the draft report against the raw KPI block for numerical "
         "inaccuracies, missing violations, and incorrect stability verdicts. Returns a numbered correction "
         "list or 'No improvements needed.' Only allowed verdicts are STABLE, CRITICALLY STABLE, UNSTABLE."),
        ("agent_6_final_report", "Agent 6", "Takes the draft + correction list and produces a polished, "
         "self-consistent final report. Applies all corrections; does not introduce new content."),
        ("agent_7_comparison", "Agent 7", "Compares KPIs and narrative across all cases. Identifies the "
         "worst-case scenario, common trends, and differentiating factors. Produces a comparative analysis "
         "report in the same structured format as Agent 6."),
        ("grid_analysis_prompt", "Agent 9 Stage A", "Analyses the structured grid_data dict to identify "
         "topological vulnerabilities: low short-circuit buses, overloaded lines, voltage deviations, "
         "and network bottlenecks."),
        ("mitigation_prompt", "Agent 9 Stage B", "Produces a ranked mitigation action plan with three "
         "sections (IMMEDIATE, SHORT-TERM, LONG-TERM). Each action entry includes: Type, Target, Issue, "
         "Mechanism, Expected Impact, and Implementation Complexity."),
    ]

    for agent_id, agent_label, desc in prompt_agents:
        p = doc.add_paragraph()
        run = p.add_run(f"{agent_label} — {agent_id}: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph(
        "All LLM calls use temperature=0.1 to ensure consistent, factual outputs. max_tokens is set "
        "conservatively per agent to avoid truncated reports while minimising API cost. The fast-tier "
        "model (MODEL_FAST) is used for Agents 4 and 5 where multiple calls may occur per case; "
        "the smart-tier model (MODEL_SMART) is reserved for final synthesis and mitigation where "
        "output quality has a direct impact on the engineering value of the report."
    )

    doc.add_page_break()

    # =========================================================
    # 9. DEPENDENCIES
    # =========================================================
    h = doc.add_heading("9. Dependencies & Installation", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)

    h2 = doc.add_heading("9.1 Required Python Packages", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    tbl_dep = doc.add_table(rows=1, cols=3)
    tbl_dep.style = 'Table Grid'
    for i, txt in enumerate(["Package", "Minimum Version", "Purpose"]):
        tbl_dep.rows[0].cells[i].text = txt
        for para in tbl_dep.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        tc = tbl_dep.rows[0].cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)

    deps = [
        ("numpy", "1.26", "Numerical KPI computation in Agent 2"),
        ("matplotlib", "3.8", "Plot generation in Agent 3 (Agg backend)"),
        ("google-genai", "1.0", "Gemini LLM provider (install if using Gemini)"),
        ("groq", "0.9", "Groq LLM provider (install if using Groq)"),
        ("python-pptx", "0.6.23", "PowerPoint generation in Agent 8 (optional, graceful degradation)"),
        ("networkx", "any", "NetworkX graph layout for topology export (optional, circular fallback)"),
        ("DIgSILENT PowerFactory API", "2025 SP1", "Simulation engine Python bindings (system install)"),
    ]
    for idx, row_d in enumerate(deps):
        row = tbl_dep.add_row()
        for i, text in enumerate(row_d):
            row.cells[i].text = text
            tc = row.cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fills[idx % 2])
            tcPr.append(shd)

    doc.add_paragraph()

    h2 = doc.add_heading("9.2 Installation Steps", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    for step, cmd in [
        ("Install core dependencies:", "pip install numpy matplotlib"),
        ("Install Gemini SDK:", "pip install google-genai"),
        ("Install Groq SDK (optional):", "pip install groq"),
        ("Install PowerPoint support (optional):", "pip install python-pptx"),
        ("Install NetworkX (optional):", "pip install networkx"),
        ("Add PowerFactory Python path:", "Set in Agent_DIgSILENT.py sys.path.insert(0, ...) block"),
    ]:
        p = doc.add_paragraph()
        p.add_run(step + "  ").bold = True
        code_run = p.add_run(cmd)
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(10)

    h2 = doc.add_heading("9.3 Configuration Setup", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    for i, step in enumerate([
        "Copy simulation_config.example.json to simulation_config.json.",
        "Set llm.provider to 'gemini' or 'groq' and enter the corresponding API key.",
        "Set project_path to the full UNC or absolute path of the PowerFactory project (.IntPrj).",
        "Set study_case and base_study_case to match case names inside the project.",
        "Set output_dir to a writable directory for all output artefacts.",
        "Optionally add a 'cases' array for multi-case batch execution.",
        "Run: python Multi_Agent_RMS_google.py",
    ], 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)

    doc.add_page_break()

    # =========================================================
    # 10. EXECUTION FLOW EXAMPLE
    # =========================================================
    h = doc.add_heading("10. End-to-End Execution Example", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)

    h2 = doc.add_heading("10.1 Single-Case Execution", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    doc.add_paragraph(
        "The following shows a complete run for a 3-phase bus fault at Bus 25 with clearing at 80 ms on a "
        "39-bus test network, using Gemini as the LLM provider:"
    )

    steps_ex = [
        ("Startup", "Multi_Agent_RMS_google.py loads simulation_config.json, configures the Gemini client "
         "(gemini-2.5-flash-lite / gemini-2.5-flash), and loads all prompts from agent_prompts.csv."),
        ("Intake", "No 'cases' array detected. Agent 0 prompts the user. User types: "
         "'3-phase bus fault at Bus 25, clearing at 80ms, simulate for 10 seconds.' "
         "LLM returns: [{\"case_name\":\"BusFault25\",\"fault_type\":\"bus\","
         "\"fault_element\":\"Bus 25.ElmTerm\",\"t_fault\":1.0,\"t_clear\":1.08,\"t_end\":10.0}]"),
        ("Simulation", "DIgSILENTAgent connects to PowerFactory 2025 SP1, activates the '39 Bus' project, "
         "activates the 'Case 1' study case (falling back to '0. Base' if not found), "
         "exports the topology PNG, runs the load flow, creates an EvtShc at Bus 25 (t=1.0 s, cleared at t=1.08 s), "
         "runs ComInc + ComSim (dt=0.01 s), exports signals to All_calculations.ElmRes, then exports to CSV."),
        ("Analysis", "agent_analysis.py parses the CSV, identifies 10 bus voltage traces and 9 generator "
         "speed/angle traces. Computes all KPIs. Bus 25 voltage nadir: 0.41 p.u. Buses 12, 25, 26 "
         "flagged as out_of_limit_after_check=True. Gen 6 speed: max=1.032 p.u., settle_s=4.2 s."),
        ("Plotting", "agent_plot.py generates voltages.png (all 10 buses), speeds.png (9 generators), "
         "and three close-up violation PNGs for Buses 12, 25, 26."),
        ("LLM Summary", "Agent 4 receives the compact KPI block and produces a 400-word draft report in "
         "~3 seconds (Gemini Flash Lite)."),
        ("LLM Review", "Agent 5 cross-checks the draft. Returns: '1. Bus 26 post-fault voltage not mentioned. "
         "2. Gen 6 settle time under-reported as 3.1s; correct value is 4.2s.'"),
        ("Final Report", "Agent 6 applies corrections, producing a complete polished report with all violations "
         "and accurate KPIs. Takes ~5 seconds (Gemini Flash)."),
        ("Presentation", "Agent 8 builds BusFault25_presentation.pptx: 1 title slide, 4 text slides, 5 image slides."),
        ("Mitigation", "Agent 9 Stage A identifies Bus 25 and Bus 26 as weak-Ikss buses. Stage B recommends: "
         "IMMEDIATE — add static VAr compensator at Bus 25; SHORT-TERM — reinforce Line 25-26; "
         "LONG-TERM — install FACTS device at Bus 12."),
        ("Artefact Save", "All outputs written to output_dir/BusFault25_20250417_143512/: "
         "8 files including the Word report, PPTX, and 5 PNGs."),
    ]
    for title, desc in steps_ex:
        p = doc.add_paragraph()
        run = p.add_run(title + ": ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        p.add_run(desc)

    h2 = doc.add_heading("10.2 Multi-Case Batch Execution", level=2)
    set_heading_color(h2, 0x26, 0x59, 0x9E)

    doc.add_paragraph(
        "When simulation_config.json contains a 'cases' array with three entries (bus fault, line trip, "
        "generator disconnection), the pipeline:"
    )
    for item in [
        "Runs each case independently through the single-case pipeline (no individual final reports unless word_document=1).",
        "After all three cases complete, calls Agent 7 (comparison_agent) with the KPI blocks and case names.",
        "Generates overlay voltage and speed comparison plots (one per violated bus, one per top-3 generator).",
        "Saves a consolidated Word document and PowerPoint with all three cases' data.",
        "Total wall-clock time: typically 2–5 minutes per case depending on simulation length and LLM latency.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    doc.add_page_break()

    # =========================================================
    # 11. ERROR HANDLING & DEGRADATION
    # =========================================================
    h = doc.add_heading("11. Error Handling and Graceful Degradation", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)

    doc.add_paragraph(
        "The pipeline distinguishes between mandatory and optional agents. Agents 1 and 2 (simulation and "
        "analysis) are mandatory: if either fails, the case is aborted and an error entry is written to the "
        "pipeline log. All other agents are optional — failures are logged as warnings and execution continues "
        "with the remaining agents."
    )

    graceful = [
        ("PowerFactory not found", "Agent 1 raises an ImportError with a helpful message indicating the "
         "expected Python path. The pipeline logs the failure and aborts the case."),
        ("python-pptx not installed", "Agent 8 returns ok=False immediately with message 'python-pptx not "
         "available'. No crash; presentation simply not generated."),
        ("NetworkX not installed", "export_grid_graph() falls back to a circular layout for the topology "
         "diagram. The grid_data dict is still fully populated."),
        ("LLM API error / timeout", "run_agent() catches exceptions and returns an empty string. The "
         "downstream agent receives an empty draft/improvement and produces a minimal output."),
        ("Study case not found", "activate_study_case() tries the target case, then falls back to "
         "base_study_case. If neither exists, returns (False, error_message)."),
        ("Generator switched out", "_is_switched_out_generator() detects gen_switch open events and "
         "excludes the affected generator from speed and angle KPIs and plots — preventing misleading "
         "stability metrics."),
        ("CSV export failure", "If ComRes returns no data, analysis_agent receives an empty dict and "
         "returns (None, None), triggering pipeline abort for that case."),
    ]
    for scenario, handling in graceful:
        p = doc.add_paragraph()
        run = p.add_run(scenario + ": ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        p.add_run(handling)

    doc.add_page_break()

    # =========================================================
    # 12. GLOSSARY
    # =========================================================
    h = doc.add_heading("12. Glossary", level=1)
    set_heading_color(h, 0x1F, 0x38, 0x64)

    terms = [
        ("DSA", "Dynamic Security Assessment — the process of evaluating whether a power system can "
         "survive a set of credible contingencies while remaining in a stable operating state."),
        ("RMS simulation", "Root Mean Square simulation — a time-domain phasor simulation method used "
         "in power systems to analyse transient behaviour on the order of milliseconds to seconds. "
         "Used by DIgSILENT ComSim."),
        ("KPI", "Key Performance Indicator — a quantitative metric extracted from simulation results "
         "to characterise system behaviour (e.g. voltage nadir, speed settling time)."),
        ("Nadir", "The lowest value of a voltage (or frequency) signal during and immediately after a fault."),
        ("Settling time", "The time after fault clearing at which a signal remains continuously within "
         "a defined band (±10% for voltage, ±0.5% for speed) for the remainder of the simulation."),
        ("p.u.", "Per-unit — a normalised unit system where nominal voltage = 1.0 p.u., "
         "nominal frequency = 1.0 p.u., etc."),
        ("Ikss", "Initial short-circuit current (kA) — a measure of the short-circuit strength at a bus. "
         "Low Ikss buses are more vulnerable to voltage collapse."),
        ("EvtShc / EvtSwitch", "PowerFactory event objects that define fault (short-circuit) or "
         "switching events in an RMS simulation scenario."),
        ("MODEL_FAST / MODEL_SMART", "Symbolic LLM tier aliases used throughout the codebase. "
         "Resolved at call time to the provider-specific model name configured in simulation_config.json."),
        ("OOXML", "Office Open XML — the XML-based file format used by Microsoft Word (.docx) and "
         "PowerPoint (.pptx). report_utils.py uses raw OOXML + stdlib zipfile to produce .docx files "
         "without a python-docx dependency."),
        ("Groq", "A hardware and API platform providing ultra-fast inference on open-source LLMs "
         "(Llama, Mixtral). Supported as an alternative to Gemini in the pipeline."),
    ]
    for term, definition in terms:
        p = doc.add_paragraph()
        run = p.add_run(term + ": ")
        run.bold = True
        p.add_run(definition)

    # =========================================================
    # SAVE
    # =========================================================
    doc.save(OUTPUT_PATH)
    print(f"\nDocumentation saved to:\n  {OUTPUT_PATH}\n")
    return OUTPUT_PATH


if __name__ == "__main__":
    build_doc()
